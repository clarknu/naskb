"""文档分析器 — 提取文件文本内容与元数据。

按架构文档 Phase 1 范围：
- PDF    → PyMuPDF (fitz) 提取文本
- Word   → python-docx 提取段落
- Excel  → openpyxl 提取表格内容摘要
- TXT/MD → chardet 编码探测（缺失时回退 utf-8/gbk 依次尝试）

所有解析器依赖均为可选：依赖未安装或解析失败时返回
ExtractionResult(text=None, error=原因)，调用方决定降级策略。
"""
from __future__ import annotations

import mimetypes
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from ..fs.base import FileSystemAdapter
from ..sidecar import Metadata
from ..exts import TEXT_EXTS  # 纯文本子集权威定义在 exts.py（B-05 统一，2026-08-24）

_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_DOC_EXTS = {".doc"}        # 老版二进制 Word（97-2003）
_XLSX_EXTS = {".xlsx", ".xlsm"}
_XLS_EXTS = {".xls"}        # 老版二进制 Excel（97-2003）
_RTF_EXTS = {".rtf"}
_KM_EXTS = {".km"}          # 思维导图 JSON 树
_MMAP_EXTS = {".mmap"}     # MindManager: ZIP + Document.xml


@dataclass
class ExtractionResult:
    """一次内容提取的结果。text 为 None 表示未能提取。"""
    text: Optional[str] = None
    metadata: Metadata = field(default_factory=Metadata)
    analyzer: str = ""
    error: Optional[str] = None
    truncated: bool = False


def _mime_type(path: str) -> str:
    mt, _ = mimetypes.guess_type(path)
    return mt or "application/octet-stream"


# ═══════════════════════════════════════════════════════════════════
# 各类型提取器
# ═══════════════════════════════════════════════════════════════════


def extract_doc(path: str, max_chars: int = 100_000,
                 com_app: Any = None) -> ExtractionResult:
    """老版二进制 .doc（Word 97-2003）文本提取。

    优先 Word COM（Windows + Microsoft Word，质量最好）；
    无 Word 时 olefile 宽松提取 WordDocument 流中的可读文本（兜底）。

    com_app: 可传入已创建的 Word.Application 实例以复用（批量分析时
    避免每个文件重新启动 Word）；为 None 时自动创建并在用完退出。
    """
    size = Path(path).stat().st_size
    text: Optional[str] = None
    error: Optional[str] = None
    owns_app = False
    try:
        if com_app is None:
            import win32com.client
            com_app = win32com.client.Dispatch("Word.Application")
            com_app.Visible = False
            com_app.DisplayAlerts = 0
            owns_app = True
        try:
            doc = com_app.Documents.Open(path, ReadOnly=True)
            text = doc.Content.Text
            doc.Close(False)
        finally:
            if owns_app:
                try:
                    com_app.Quit()
                except Exception:
                    pass
    except Exception as e:
        error = f"Word COM 不可用: {e}"
    if not text:
        try:
            text = _extract_doc_ole(path)
        except Exception as e:
            error = f"DOC 解析失败: {e}"
    truncated = bool(text and len(text) > max_chars)
    if truncated:
        text = text[:max_chars]
    return ExtractionResult(
        text=text or None,
        metadata=Metadata(file_type="application/msword", file_size=size),
        analyzer="doc",
        error=None if text else (error or "DOC 无文本"),
        truncated=truncated,
    )


def _extract_doc_ole(path: str) -> Optional[str]:
    """兜底：无 Word 时用 olefile 从 WordDocument 流宽松提取可读文本。

    将流按 UTF-16LE 解释，抽取连续可读片段（中文/ASCII/全角）。
    质量低于 Word COM，但保证无 Office 环境也能拿到正文。
    """
    import olefile
    import re

    ole = olefile.OleFileIO(path)
    try:
        if not ole.exists("WordDocument"):
            return None
        data = ole.openstream("WordDocument").read()
    finally:
        ole.close()
    text = data.decode("utf-16-le", errors="ignore")
    parts = [m.group() for m in re.finditer(
        r"[\u4e00-\u9fff\u0020-\u007e\uff00-\uffef]{4,}", text)]
    if not parts:
        return None
    return "\n".join(parts)


def extract_xls(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """老版二进制 .xls（Excel 97-2003）→ xlrd 提取工作表内容。"""
    import xlrd

    size = Path(path).stat().st_size
    try:
        book = xlrd.open_workbook(path)
        lines: list[str] = []
        for sh in book.sheets():
            lines.append(f"## Sheet: {sh.name}")
            for r in range(min(sh.nrows, 1000)):
                vals = [str(sh.cell_value(r, c)).strip()
                        for c in range(sh.ncols)]
                line = " | ".join(v for v in vals if v)
                if line:
                    lines.append(line)
        text = "\n".join(lines)
        truncated = len(text) > max_chars
        if truncated:
            text = text[:max_chars]
        return ExtractionResult(
            text=text or None,
            metadata=Metadata(file_type="application/vnd.ms-excel", file_size=size),
            analyzer="xls",
            error=None if text else "XLS 无内容",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractionResult(
            metadata=Metadata(file_type="application/vnd.ms-excel", file_size=size),
            analyzer="xls",
            error=f"XLS 解析失败（应为核心 BIFF/OLE 格式）: {e}",
        )


def extract_text_file(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """TXT/MD/CSV 等纯文本文件，编码探测后读取。"""
    data = Path(path).read_bytes()
    text = _decode(data)
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars]
    size = Path(path).stat().st_size
    return ExtractionResult(
        text=text,
        metadata=Metadata(file_type=_mime_type(path), file_size=size),
        analyzer="text",
        truncated=truncated,
    )


def extract_rtf(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """RTF 富文本：轻量剥离控制字提取纯文本（不依赖外部库）。

    中文内容通常以 \\uN 转义 + 实际字符存储，粗提取可保留大部分文本。
    """
    import re

    size = Path(path).stat().st_size
    try:
        text = Path(path).read_bytes().decode("latin-1", errors="replace")
    except OSError as e:
        return ExtractionResult(
            metadata=Metadata(file_type="application/rtf", file_size=size),
            analyzer="rtf", error=f"RTF 读取失败: {e}",
        )
    text = re.sub(r"\\par[ds]?", "\n", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)   # \'xx 字节转义（latin-1 已覆盖）
    text = re.sub(r"\\u-?\d+ ?", "", text)          # \uN 转义（数字删掉，字符保留）
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)  # 其他控制字
    text = re.sub(r"[{}]", "", text)
    text = text.replace("\\", "")
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    out = "\n".join(lines)
    truncated = len(out) > max_chars
    if truncated:
        out = out[:max_chars]
    return ExtractionResult(
        text=out or None,
        metadata=Metadata(file_type="application/rtf", file_size=size),
        analyzer="rtf",
        error=None if out else "RTF 未提取到可读文本",
        truncated=truncated,
    )


def extract_km(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """KM（思维导图 JSON 树）→ 树形文本（保留层级缩进）。

    结构: {"root": {"data": {"text": ...}, "children": [...]}}
    """
    import json

    size = Path(path).stat().st_size
    try:
        raw = _decode(Path(path).read_bytes())
        tree = json.loads(raw)
    except (json.JSONDecodeError, ValueError, OSError) as e:
        return ExtractionResult(
            metadata=Metadata(file_type="application/json", file_size=size),
            analyzer="km",
            error=f"KM 解析失败（应为核心 JSON 文本）: {e}",
        )

    lines: list[str] = []

    def walk(node: Any, depth: int) -> None:
        if not isinstance(node, dict):
            return
        d = node.get("data") or {}
        t = str(d.get("text", "")).strip()
        if t:
            lines.append("  " * depth + t)
        for child in node.get("children") or []:
            walk(child, depth + 1)

    walk(tree.get("root") or tree, 0)
    out = "\n".join(lines)
    truncated = len(out) > max_chars
    if truncated:
        out = out[:max_chars]
    return ExtractionResult(
        text=out or None,
        metadata=Metadata(file_type="application/json", file_size=size),
        analyzer="km",
        error=None if out else "KM 文件未包含可读的节点文本",
        truncated=truncated,
    )


def extract_mmap(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """MMAP（MindManager ZIP）→ Document.xml 中 ap:Text 的 PlainText 属性文本。"""
    import html as html_mod
    import re
    import zipfile

    size = Path(path).stat().st_size
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("Document.xml").decode("utf-8", errors="replace")
    except Exception as e:
        return ExtractionResult(
            metadata=Metadata(file_type="application/x-mindmanager", file_size=size),
            analyzer="mmap",
            error=f"MMAP 打开失败（应为核心 ZIP 容器）: {e}",
        )
    texts = re.findall(r'<ap:Text[^>]*PlainText="([^"]*)"', xml)
    out = "\n".join(html_mod.unescape(t) for t in texts)
    truncated = len(out) > max_chars
    if truncated:
        out = out[:max_chars]
    return ExtractionResult(
        text=out or None,
        metadata=Metadata(file_type="application/x-mindmanager", file_size=size),
        analyzer="mmap",
        error=None if out else "MMAP 未包含可读文本（无 PlainText 节点）",
        truncated=truncated,
    )


def _decode(data: bytes) -> str:
    """按 chardet → utf-8 → gbk → latin-1 顺序尝试解码。"""
    try:
        import chardet
        guess = chardet.detect(data)
        encoding = guess.get("encoding") or "utf-8"
        return data.decode(encoding, errors="replace")
    except ImportError:
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return data.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
    return data.decode("utf-8", errors="replace")


def extract_pdf(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """PDF → PyMuPDF 提取文本。"""
    size = Path(path).stat().st_size
    try:
        import pymupdf as fitz  # PyMuPDF >= 1.24 新入口
    except ImportError:
        try:
            import fitz  # 旧版兼容 (PyMuPDF < 1.24)
        except ImportError:
            return ExtractionResult(
                metadata=Metadata(file_type="application/pdf", file_size=size),
                analyzer="pdf",
                error="PyMuPDF 未安装 (pip install pymupdf)",
            )
    try:
        doc = fitz.open(path)
        parts: list[str] = []
        page_count = len(doc)
        img_pages = 0
        for page in doc:
            parts.append(page.get_text())
            try:
                if page.get_images():
                    img_pages += 1
            except Exception:
                pass
        text = "\n".join(parts).strip()
        # 扫描件特征：存在图片页 且 文本稀薄（平均每页 < 2000 字符）。
        # 这类 PDF 的文本层常是劣质 OCR 结果（错字多），应交给 MinerU 重新识别。
        scan_like = bool(img_pages > 0 and page_count > 0
                         and (len(text) / page_count) < 2000)
        truncated = len(text) > max_chars
        if truncated:
            text = text[:max_chars]
        width = doc[0].rect.width if len(doc) > 0 else None
        height = doc[0].rect.height if len(doc) > 0 else None
        doc.close()
        return ExtractionResult(
            text=text or None,
            metadata=Metadata(file_type="application/pdf", file_size=size,
                              width=width, height=height,
                              exif={"scan_like": scan_like,
                                    "page_count": page_count}),
            analyzer="pdf",
            error=None if text else "PDF 无文本层（扫描件需 OCR，属 Phase 2）",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractionResult(
            metadata=Metadata(file_type="application/pdf", file_size=size),
            analyzer="pdf",
            error=f"PDF 解析失败: {e}",
        )


def extract_docx(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """Word → python-docx 提取段落与表格。"""
    size = Path(path).stat().st_size
    try:
        import docx
    except ImportError:
        return ExtractionResult(
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                file_size=size),
            analyzer="docx",
            error="python-docx 未安装 (pip install python-docx)",
        )
    try:
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        truncated = len(text) > max_chars
        if truncated:
            text = text[:max_chars]
        return ExtractionResult(
            text=text or None,
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                file_size=size),
            analyzer="docx",
            error=None if text else "Word 文档无文本内容",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractionResult(
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                file_size=size),
            analyzer="docx",
            error=f"Word 解析失败: {e}",
        )


def extract_xlsx(path: str, max_chars: int = 100_000) -> ExtractionResult:
    """Excel → openpyxl 提取单元格摘要。"""
    size = Path(path).stat().st_size
    try:
        import openpyxl
    except ImportError:
        return ExtractionResult(
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                file_size=size),
            analyzer="xlsx",
            error="openpyxl 未安装 (pip install openpyxl)",
        )
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"## Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
            if len("\n".join(parts)) > max_chars:
                break
        text = "\n".join(parts).strip()
        truncated = len(text) > max_chars
        if truncated:
            text = text[:max_chars]
        wb.close()
        return ExtractionResult(
            text=text or None,
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                file_size=size),
            analyzer="xlsx",
            error=None if text else "Excel 无单元格内容",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractionResult(
            metadata=Metadata(file_type=(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                file_size=size),
            analyzer="xlsx",
            error=f"Excel 解析失败: {e}",
        )


# ═══════════════════════════════════════════════════════════════════
# 分发器
# ═══════════════════════════════════════════════════════════════════


class DocumentAnalyzer:
    """按文件扩展名分发到对应提取器。"""

    def __init__(self, max_chars: int = 100_000,
                 max_file_bytes: Optional[int] = None,
                 com_app: Any = None):
        """
        com_app: 复用的 Word.Application 实例（批量分析 .doc 时传入，
        避免每个文件启动一次 Word）；由调用方负责关闭。
        """
        self.max_chars = max_chars
        self.max_file_bytes = max_file_bytes
        self._com_app = com_app

    def close_com(self) -> None:
        """关闭批量复用期间创建的 Word 实例。"""
        if self._com_app is not None:
            try:
                self._com_app.Quit()
            except Exception:
                pass
            self._com_app = None

    def extract(self, path: str) -> ExtractionResult:
        """提取本地文件内容。"""
        ext = Path(path).suffix.lower()
        if ext in _PDF_EXTS:
            return extract_pdf(path, self.max_chars)
        if ext in _DOCX_EXTS:
            return extract_docx(path, self.max_chars)
        if ext in _DOC_EXTS:
            return extract_doc(path, self.max_chars, com_app=self._com_app)
        if ext in _XLSX_EXTS:
            return extract_xlsx(path, self.max_chars)
        if ext in _XLS_EXTS:
            return extract_xls(path, self.max_chars)
        if ext in _RTF_EXTS:
            return extract_rtf(path, self.max_chars)
        if ext in _KM_EXTS:
            return extract_km(path, self.max_chars)
        if ext in _MMAP_EXTS:
            return extract_mmap(path, self.max_chars)
        if ext in TEXT_EXTS:
            return extract_text_file(path, self.max_chars)
        return ExtractionResult(
            metadata=Metadata(file_type=_mime_type(path),
                              file_size=Path(path).stat().st_size),
            analyzer="unknown",
            error=f"不支持的文件类型: {ext or '(无扩展名)'}",
        )

    def extract_remote(self, fs: FileSystemAdapter, remote_path: str,
                       tmp_dir: str) -> ExtractionResult:
        """下载远端文件到临时目录后提取，处理完立即删除临时文件。

        大文件策略（架构文档 6.4）：超过 max_file_bytes 的文件直接拒绝，
        不下载；小文件下载后分析。
        """
        import os

        os.makedirs(tmp_dir, exist_ok=True)
        # 大文件过滤：先 stat 再决定是否下载
        st = fs.stat(remote_path)
        if st and self.max_file_bytes is not None and st.size_bytes > self.max_file_bytes:
            return ExtractionResult(
                metadata=Metadata(file_type=_mime_type(remote_path),
                                  file_size=st.size_bytes),
                analyzer="remote",
                error=(f"文件过大 ({st.size_bytes} bytes > "
                       f"{self.max_file_bytes} bytes)，已跳过下载"),
            )
        name = Path(remote_path).name
        local = os.path.join(tmp_dir, f"{os.getpid()}-{name}")
        try:
            with open(local, "wb") as f:
                for chunk in fs.read_chunks(remote_path):
                    f.write(chunk)
            result = self.extract(local)
            # 补全远端元数据
            if st:
                result.metadata.file_size = st.size_bytes
            return result
        finally:
            try:
                os.remove(local)
            except OSError:
                pass
