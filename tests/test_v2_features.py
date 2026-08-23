"""V2 特性测试：adopt/export、解析视图、Office 简版预览、缩略图、pg-rebind。

PG 门控部分与 test_content_access 同策略（[pg] 不可达即跳过）。
"""
import hashlib
import os

import pytest

try:
    import psycopg  # noqa: F401
    import pgvector  # noqa: F401
    from fastapi.testclient import TestClient
    _DEPS = True
except ImportError:
    _DEPS = False

pytestmark = pytest.mark.skipif(not _DEPS, reason="psycopg/pgvector/httpx 未安装")


class _RealCfg:
    def __init__(self):
        from naskb.common.config import Config
        self._c = Config.from_work_path("NASKB_data")
        for k in ("pg_host", "pg_port", "pg_user", "pg_password",
                  "pg_database", "work_path", "server_tokens",
                  "anonymous_read"):
            setattr(self, k, getattr(self._c, k))

    @property
    def pg_enabled(self):
        return bool(self.pg_host)


@pytest.fixture(scope="module")
def env(tmp_path_factory):
    try:
        cfg = _RealCfg()
    except Exception:
        pytest.skip("config.toml 不可用")
    from naskb.common.pgstore import PgStore
    pg = PgStore(cfg)
    try:
        with pg.connect() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT 1")
    except Exception:
        pytest.skip("[pg] 不可达")
    return cfg, pg


class _FakeEmbedder:
    """零向量假嵌入器（adopt 测试用，避免下载模型）。"""

    def encode(self, texts):
        import numpy as np
        from naskb.common.pgstore import EMBEDDING_DIM
        return np.zeros((len(texts), EMBEDDING_DIM), dtype="float32")

    def encode_one(self, text):
        import numpy as np
        from naskb.common.pgstore import EMBEDDING_DIM
        return np.zeros((EMBEDDING_DIM,), dtype="float32")

    def close(self):
        pass


def _make_repo(root, files, artifacts=None):
    """在 root 下构造一个 .naskb 描述仓库（index.json + folder.json）。"""
    from naskb.common.desc_store import FileEntry, NaskbStore
    from naskb.common.fs.local import LocalAdapter
    store = NaskbStore(LocalAdapter(str(root)))
    artifacts = artifacts or {}
    for rel, content in files.items():
        entry = FileEntry(
            path=os.path.basename(rel), summary=content,
            category="文档", tags=["测试"], file_type="txt",
            size_bytes=len(content), file_hash=hashlib.sha1(
                content.encode()).hexdigest(), hash_algorithm="sha256:full",
        )
        if rel in artifacts:
            entry.exif["mineru_artifacts"] = artifacts[rel]
        store.set_entry(rel, entry)
    return store


class TestAdoptExport:
    def test_adopt_imports_repo(self, env, tmp_path):
        cfg, pg = env
        from naskb.common.source_registry import SourceRecord
        root = tmp_path / "repo"
        (root / "docs").mkdir(parents=True)
        _make_repo(root, {"docs/a.txt": "内容甲", "docs/b.txt": "内容乙"})
        rec = SourceRecord(alias="adopt-t", protocol="local",
                           root_path=str(root), access_mode="rw")
        rec.schema_name = "nas_local_local_0_uanon"

        from naskb.common.adopt import adopt_repo
        res = adopt_repo(pg, cfg, rec, embedder=_FakeEmbedder())
        assert res["adopted"] == 2
        rows = {r["rel_path"]: r for r in pg.all_rows(
            rec.schema_name, rec.source_id)}
        assert "docs/a.txt" in rows and rows["docs/a.txt"]["summary"] == "内容甲"
        dirs, files = pg.list_dir(rec.schema_name, rec.source_id, "docs")
        assert len(files) == 2
        # 目录描述卡片已生成
        ddirs, _ = pg.list_dir(rec.schema_name, rec.source_id, "")
        assert any(d["name"] == "docs" for d in ddirs)
        pg.delete_source_rows(rec.schema_name, rec.source_id)

    def test_export_rebuilds_repo(self, env, tmp_path):
        cfg, pg = env
        import secrets
        sid = __import__("uuid").uuid4()
        schema = "nas_test_" + secrets.token_hex(4)
        pg.ensure_nas_schema(schema)
        pg.reconcile_resources(schema, sid, [
            {"rel_path": "docs/hello.txt", "name": "hello.txt",
             "parent_dir": "docs", "size_bytes": 11,
             "mtime": 0.0, "ctime": 0.0, "file_hash": "h1",
             "hash_algorithm": "sha256:full", "file_type": "txt"}])
        from naskb.common.source_registry import SourceRecord
        rec = SourceRecord(alias="export-t", protocol="local",
                           root_path=str(tmp_path), access_mode="rw")
        rec.schema_name = schema
        rec.source_id = str(sid)
        out = tmp_path / "out"
        from naskb.common.adopt import export_repo
        res = export_repo(pg, cfg, rec, str(out))
        assert res["exported"] == 1
        idx = out / "docs" / ".naskb" / "index.json"
        assert idx.is_file()
        assert "hello.txt" in idx.read_text(encoding="utf-8")
        pg.delete_source_rows(schema, sid)


class TestParsedView:
    """解析视图：adopt 时登记 artifacts → /parsed 从源端流式出 HTML。"""

    def test_parsed_endpoint(self, env, tmp_path):
        cfg, pg = env
        import uuid as _uuid
        from naskb.server.app import create_app

        root = tmp_path / "repo"
        (root / "docs").mkdir(parents=True)
        (root / "docs" / "x.pdf").write_bytes(b"%PDF-1.4 fake")
        (root / "docs" / ".naskb" / "artifacts" / "x").mkdir(parents=True)
        (root / "docs" / ".naskb" / "artifacts" / "x" / "out.html").write_text(
            "<html><body>解析后的版面</body></html>", encoding="utf-8")
        art = {"html_path": "artifacts/x/out.html"}
        _make_repo(root, {"docs/x.pdf": "扫描件摘要"},
                   artifacts={"docs/x.pdf": art})

        alias = "parsed-" + _uuid.uuid4().hex[:6]
        app = create_app(cfg)
        with TestClient(app) as c:
            d = c.post("/api/sources", json={
                "alias": alias, "protocol": "local",
                "root_path": str(root), "access_mode": "rw"})
            sid = d.json()["source"]["source_id"]
            try:
                rec = c.app.state.registry.get(sid)
                from naskb.common.adopt import adopt_repo
                adopt_repo(pg, cfg, rec, embedder=_FakeEmbedder())
                rows = pg.all_rows(rec.schema_name, rec.source_id)
                target = [r for r in rows
                          if r["rel_path"] == "docs/x.pdf"][0]
                assert target["artifacts"].get("html_path")
                r = c.get(f"/api/files/{target['resource_id']}/parsed",
                          params={"src": sid})
                assert r.status_code == 200, r.text
                assert "解析后的版面" in r.text
            finally:
                c.delete(f"/api/sources/{sid}?purge=true")


class TestOfficeAndThumb:
    def test_docx_to_html(self, tmp_path):
        import docx
        from naskb.server.office import docx_to_html
        f = tmp_path / "a.docx"
        d = docx.Document()
        d.add_heading("标题一", level=1)
        d.add_paragraph("正文段落内容")
        d.add_table(rows=2, cols=2).rows[0].cells[0].text = "甲"
        d.save(f)
        html = docx_to_html(str(f))
        assert html and "标题一" in html and "正文段落内容" in html

    def test_xlsx_to_html(self, tmp_path):
        from openpyxl import Workbook
        from naskb.server.office import xlsx_to_html
        f = tmp_path / "b.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(["列一", "列二"])
        ws.append(["v1", "v2"])
        wb.save(f)
        html = xlsx_to_html(str(f))
        assert html and "列一" in html and "v1" in html

    def test_thumbnail_image_cached(self, env, tmp_path, tmp_path_factory):
        cfg, pg = env
        import secrets
        from PIL import Image
        from naskb.common.fs.local import LocalAdapter
        sid = __import__("uuid").uuid4()
        schema = "nas_test_" + secrets.token_hex(4)
        pg.ensure_nas_schema(schema)
        root = tmp_path_factory.mktemp("thumb")
        img = Image.new("RGB", (1000, 800), (200, 60, 60))
        img.save(root / "p.png")
        pg.reconcile_resources(schema, sid, [
            {"rel_path": "p.png", "name": "p.png", "parent_dir": "",
             "size_bytes": os.path.getsize(root / "p.png"),
             "mtime": 0.0, "ctime": 0.0, "file_hash": "",
             "hash_algorithm": "", "file_type": "png"}])
        from naskb.server.thumb import thumbnail
        row = pg.get_resource(schema, _rid(pg, schema, sid, "p.png"))
        data = thumbnail(pg, cfg, row, LocalAdapter(str(root)), w=320)
        assert data and data[:2] == b"\xff\xd8"     # JPEG magic
        # 二次命中缓存（不再解码）
        cache = os.path.join(cfg.work_path, "store", "thumbs")
        assert any(f.endswith(".jpg") for f in os.listdir(cache))
        pg.delete_source_rows(schema, sid)

    def test_rebind_nas(self, env):
        from naskb.common.pgstore import schema_name_for
        cfg, pg = env
        import secrets
        host = "10.0.0." + str(secrets.randbelow(200) + 1)
        schema_old = schema_name_for("webdav", host, 5006, "u1")
        pg.ensure_nas_schema(schema_old)
        pg.get_or_create_nas("webdav", host, 5006, "u1")
        res = pg.rebind_nas({"protocol": "webdav", "host": host,
                             "port": 5006, "username": "u1"},
                            {"protocol": "webdav", "host": host,
                             "port": 6006, "username": "u1"})
        assert res["changed"] is True
        assert res["new_schema"] != res["old_schema"]
        # 旧 schema 已改名，新 schema 可用
        from naskb.common.pgstore import LEGACY_SOURCE
        pg.reconcile_resources(res["new_schema"], LEGACY_SOURCE, [])
        # 清理
        import psycopg
        with pg.connect() as conn:
            with conn.cursor() as cur:
                cur.execute("DROP SCHEMA IF EXISTS %s CASCADE"
                            % schema_name_for("webdav", host, 6006, "u1"))


def _rid(pg, schema, sid, rel):
    parent = rel.rsplit("/", 1)[0] if "/" in rel else ""
    _, files = pg.list_dir(schema, sid, parent)
    return next(f["resource_id"] for f in files if f["rel_path"] == rel)
