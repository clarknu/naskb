"""批量分析引擎测试：并发、增量幂等、类型分发、COM 复用参数。"""
import json
import threading
import time

import pytest

from naskb.common.batch import analyze_tree
from naskb.common.config import Config
from naskb.common.desc_store import NaskbStore
from naskb.common.fs.local import LocalAdapter
from naskb.common.analyzer.image import IMAGE_STRUCTURE_PROMPT


@pytest.fixture
def work(tmp_path):
    cfg = Config.from_work_path(str(tmp_path))
    return cfg


class FakeLLM:
    """mock LLM：记录并发数，返回固定摘要。"""

    def __init__(self, delay: float = 0.05):
        self.delay = delay
        self.active = 0
        self.max_active = 0
        self.calls = 0
        self._lock = threading.Lock()

    def complete_json(self, prompt: str) -> dict:
        with self._lock:
            self.active += 1
            self.max_active = max(self.max_active, self.active)
            self.calls += 1
        time.sleep(self.delay)
        with self._lock:
            self.active -= 1
        return {"summary": "测试摘要", "tags": ["测试"], "category": "文档",
                "confidence": 0.9}

    def complete(self, prompt: str) -> str:
        return "ok"

    def complete_image(self, image_path: str, prompt: str,
                       system=None) -> str:
        self.image_prompts = getattr(self, "image_prompts", [])
        self.image_prompts.append(prompt)
        return "这是嵌入扫描图: 身份证正反面"

    def close(self):
        pass


def _make_files(tmp_path, n=6):
    for i in range(n):
        (tmp_path / f"笔记{i}.md").write_text(
            f"# 笔记 {i}\n这是第 {i} 篇测试笔记内容，用于批量分析。",
            encoding="utf-8")
    (tmp_path / "数据.xlsx").write_bytes(b"PK\x03\x04 fake")
    (tmp_path / "未知.xyz").write_bytes(b"??")
    return tmp_path


class TestBatchDocs:
    def test_no_llm_analyzes_and_writes(self, tmp_path, work):
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        r = analyze_tree(fs, store, work, ".", llm=False)
        assert r.analyzed == 7          # 6 md + 1 xlsx(fake 提取失败也算 analyzed)
        assert r.unsupported == 1       # .xyz
        assert r.ignored_recorded == 1  # .xyz 仅记录名称推断
        # 条目写入（独立原数据文件）
        e = store.get_entry(str(tmp_path / "笔记0.md"))
        assert e is not None
        assert "第 0 篇" in (e.ocr_text or "")
        assert e.summary  # 无 LLM 时取提取文本前 200 字
        # 忽略文件条目：仅名称推断，未分析内容
        e2 = store.get_entry(str(tmp_path / "未知.xyz"))
        assert e2 is not None
        assert e2.processing_policy == "metadata_only"
        assert "未知.xyz" in e2.summary
        assert not e2.ocr_text and not e2.content_description

    def test_incremental_skip(self, tmp_path, work):
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        analyze_tree(fs, store, work, ".", llm=False)
        r2 = analyze_tree(fs, store, work, ".", llm=False)
        # 6 个 md 有内容被跳过；fake xlsx 提取无内容，重新尝试；
        # .xyz 忽略条目已有 → 增量跳过
        assert r2.skipped == 7
        assert r2.analyzed == 1
        assert r2.ignored_recorded == 0

    def test_hidden_and_excluded_dirs_skipped(self, tmp_path, work):
        """隐藏目录（.git）完全跳过；配置排除目录（node_modules）不记录
        文件条目但生成目录级 folder.json；系统垃圾/锁文件跳过。"""
        (tmp_path / ".git").mkdir()
        (tmp_path / ".git" / "config").write_text("repo", encoding="utf-8")
        (tmp_path / ".git" / "objects").mkdir()
        (tmp_path / ".git" / "objects" / "abc123").write_bytes(b"\x00\x01")
        (tmp_path / "node_modules").mkdir()
        (tmp_path / "node_modules" / "lib.js").write_text("x", encoding="utf-8")
        (tmp_path / "scripts").mkdir()
        (tmp_path / "scripts" / "build.js").write_text("x", encoding="utf-8")
        (tmp_path / "scripts" / "run.sh").write_text("x", encoding="utf-8")
        (tmp_path / "a.txt").write_text("内容", encoding="utf-8")
        # 系统垃圾文件 / Office 锁文件：完全跳过
        (tmp_path / "Thumbs.db").write_bytes(b"\x00")
        (tmp_path / ".DS_Store").write_bytes(b"\x00")
        (tmp_path / "~$a.docx").write_bytes(b"\x00")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        r = analyze_tree(fs, store, work, ".", llm=False)
        assert r.supported == 1
        assert r.unsupported == 2       # scripts 下 2 个计入；node_modules 不计
        assert r.ignored_recorded == 2  # build.js / run.sh 仅记录名称推断
        # 被忽略目录（scripts 全是脚本）→ folder.json 目录级描述
        fe = store.read_folder(str(tmp_path / "scripts"))
        assert fe is not None
        assert fe.file_type_distribution.get(".js") == 1
        assert fe.file_type_distribution.get(".sh") == 1
        # 排除目录（node_modules）→ 目录级 folder.json，但不记录文件条目
        fe = store.read_folder(str(tmp_path / "node_modules"))
        assert fe is not None
        assert fe.file_type_distribution.get(".js") == 1
        assert store.get_entry(str(tmp_path / "node_modules" / "lib.js")) is None
        # 隐藏目录不生成描述仓库；垃圾/锁文件无条目
        assert store.read_folder(str(tmp_path / ".git")) is None
        assert store.get_entry(str(tmp_path / ".git" / "config")) is None
        assert store.get_entry(str(tmp_path / "Thumbs.db")) is None
        assert store.get_entry(str(tmp_path / "~$a.docx")) is None

    def test_big_files_recorded_not_analyzed(self, tmp_path, work):
        """超过 max_file_mb 的支持类型文件：不下载不分析，仅记录"文件过大"。"""
        work.analyzer_max_file_mb = 1  # 1MB 上限
        big = tmp_path / "大照片.jpg"
        big.write_bytes(b"\xff\xd8" + b"x" * (2 * 1024 * 1024))
        (tmp_path / "小文档.md").write_text("小文档", encoding="utf-8")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        r = analyze_tree(fs, store, work, ".", llm=False)
        assert r.supported == 1         # 小文档进入分析
        assert r.ignored_recorded == 1  # 大照片仅记录
        e = store.get_entry(str(big))
        assert e is not None
        assert e.processing_policy == "metadata_only"
        assert "未下载分析" in e.summary
        assert "1MB" in e.summary

    def test_well_known_name_inference(self, tmp_path, work):
        """无扩展名/点开头知名文件按文件名推断意义。"""
        (tmp_path / "Dockerfile").write_text("FROM python", encoding="utf-8")
        (tmp_path / ".env").write_text("KEY=1", encoding="utf-8")
        (tmp_path / "go.mod").write_text("module x", encoding="utf-8")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        r = analyze_tree(fs, store, work, ".", llm=False)
        assert r.ignored_recorded == 3
        e = store.get_entry(str(tmp_path / "Dockerfile"))
        assert "Docker 容器构建文件" in e.summary
        e2 = store.get_entry(str(tmp_path / ".env"))
        assert "环境变量配置文件" in e2.summary
        e3 = store.get_entry(str(tmp_path / "go.mod"))
        assert "Go 模块定义" in e3.summary

    def test_deleted_file_cleans_orphan_and_cascades(self, tmp_path, work):
        """文件删除后：重跑 analyze-tree 自动清孤儿条目，目录及其祖先
        folder.json 重算（统计不再含已删文件）。"""
        (tmp_path / "a.txt").write_text("a", encoding="utf-8")
        (tmp_path / "b.txt").write_text("b", encoding="utf-8")
        (tmp_path / "子目录").mkdir()
        (tmp_path / "子目录" / "c.txt").write_text("c", encoding="utf-8")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        analyze_tree(fs, store, work, ".", llm=False)
        assert store.get_entry(str(tmp_path / "a.txt")) is not None
        fe = store.read_folder(str(tmp_path))
        assert fe.file_type_distribution.get(".txt") == 3  # 递归统计含子目录

        # 删除 a.txt 后重跑：条目清理 + 根目录 folder.json 重算
        (tmp_path / "a.txt").unlink()
        r2 = analyze_tree(fs, store, work, ".", llm=False)
        assert r2.orphans_removed == 1
        assert store.get_entry(str(tmp_path / "a.txt")) is None
        assert store.get_entry(str(tmp_path / "b.txt")) is not None
        fe2 = store.read_folder(str(tmp_path))
        assert fe2.file_type_distribution.get(".txt") == 2
        # 已删文件的独立原数据文件也清理
        assert not (tmp_path / ".naskb" / "files" / "a.txt.json").exists()

    def test_force_reanalyzes(self, tmp_path, work):
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        analyze_tree(fs, store, work, ".", llm=False)
        r2 = analyze_tree(fs, store, work, ".", llm=False, force=True)
        assert r2.analyzed == 7

    def test_limit(self, tmp_path, work):
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        r = analyze_tree(fs, store, work, ".", llm=False, limit=3)
        assert r.analyzed == 3

    def test_llm_concurrency_capped(self, tmp_path, work):
        """DeepSeek 并发数受 workers 限制。"""
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        fake = FakeLLM(delay=0.1)
        r = analyze_tree(fs, store, work, ".", llm=True, workers=2,
                         clients={"text": fake, "vision": fake, "audio": fake})
        assert fake.max_active <= 2
        assert fake.calls == 7  # 6 个 md 走 LLM + 1 次目录级 folder.json（级联）
        # LLM 摘要已合并写入
        e = store.get_entry(str(tmp_path / "笔记0.md"))
        assert e.summary == "测试摘要"
        assert e.category == "文档"
        assert e.tags == ["测试"]
        # 目录级描述已生成（受影响目录 + 祖先，同一 mock 返回固定摘要）
        fe = store.read_folder(str(tmp_path))
        assert fe is not None
        assert fe.summary == "测试摘要"

    def test_folder_cascade_to_ancestors(self, tmp_path, work):
        """新增文件后，受影响目录及其祖先的 folder.json 逐级更新。"""
        (tmp_path / "a" / "b").mkdir(parents=True)
        (tmp_path / "a" / "b" / "照片.jpg").write_bytes(b"\xff\xd8 fake")
        (tmp_path / "a" / "b" / "脚本.py").write_text("print(1)", encoding="utf-8")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)

        class _FakeFolderLLM:
            def complete_json(self, prompt):
                return {"description": "目录描述", "structure": [],
                        "tags": ["目录"], "summary": "目录摘要",
                        "language": "zh", "confidence": 0.8}

            def complete_image(self, image_path, prompt, system=None):
                return "一张测试图片"

            def close(self):
                pass

        fake = _FakeFolderLLM()
        r = analyze_tree(fs, store, work, ".",
                         llm=True, workers=2,
                         clients={"text": fake, "vision": fake, "audio": fake})
        # 受影响目录 b + 祖先 a + 根，共 3 个目录的 folder.json
        assert r.folder_updated == 3
        for d in (str(tmp_path / "a" / "b"), str(tmp_path / "a"), str(tmp_path)):
            fe = store.read_folder(d)
            assert fe is not None, d
            assert fe.description == "目录描述"
        # 忽略文件条目（脚本.py 仅记录名称推断）
        e = store.get_entry(str(tmp_path / "a" / "b" / "脚本.py"))
        assert e is not None
        assert "Python 源代码" in e.summary

    def test_subdirs_and_repo_exclusion(self, tmp_path, work):
        """子目录文件也分析；.naskb 仓库内部文件被排除。"""
        (tmp_path / "子目录").mkdir()
        (tmp_path / "子目录" / "a.txt").write_text("子目录内容", encoding="utf-8")
        (tmp_path / "b.txt").write_text("根内容", encoding="utf-8")
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        analyze_tree(fs, store, work, ".", llm=False)
        # 再跑一次不会把 .naskb 内部文件当新文件
        r2 = analyze_tree(fs, store, work, ".", llm=False)
        assert r2.skipped == 2

    def test_image_only_docx_uses_vision(self, tmp_path, work):
        """图片型 docx（无文本层）→ 提取嵌入图走视觉识别。"""
        import io
        from PIL import Image as PILImage
        import docx as docxlib

        p = tmp_path / "扫描件.docx"
        d = docxlib.Document()
        buf = io.BytesIO()
        PILImage.new("RGB", (64, 64), color=(200, 200, 255)).save(buf, "JPEG")
        buf.seek(0)
        d.add_picture(buf)
        d.save(str(p))

        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        fake = FakeLLM()
        r = analyze_tree(fs, store, work, ".", llm=True, workers=2,
                         clients={"text": fake, "vision": fake, "audio": fake})
        e = store.get_entry(str(p))
        assert e is not None
        # summary 会被 DeepSeek 摘要（fake 固定值）覆盖，识别文本在 ocr_text
        assert "嵌入扫描图" in (e.ocr_text or "")
        assert "嵌入扫描图" in (e.content_description or "")
        # 档位 1：图文流带顺序、锚点、媒体名；识别用的是结构 prompt
        assert "[图片 1]" in (e.ocr_text or "")
        assert "image1" in (e.ocr_text or "")
        assert "段落" in (e.ocr_text or "")
        assert fake.image_prompts and all(
            pr == IMAGE_STRUCTURE_PROMPT for pr in fake.image_prompts)
        assert e.category  # DeepSeek 摘要阶段会基于识别文本重新分类
        assert r.failed == 0

    def test_docx_with_text_keeps_flow(self, tmp_path, work):
        """含文本层 + 图片的 docx：正文进 ocr_text，图文流进 content_description。"""
        import io
        from PIL import Image as PILImage
        import docx as docxlib

        p = tmp_path / "图文.docx"
        d = docxlib.Document()
        d.add_paragraph("这是正文第一段")
        buf = io.BytesIO()
        PILImage.new("RGB", (64, 64), color=(200, 200, 255)).save(buf, "JPEG")
        buf.seek(0)
        d.add_picture(buf)
        d.add_paragraph("这是正文第二段")
        d.save(str(p))

        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)
        fake = FakeLLM()
        r = analyze_tree(fs, store, work, ".", llm=True, workers=2,
                         clients={"text": fake, "vision": fake, "audio": fake})
        e = store.get_entry(str(p))
        assert e is not None
        # 正文完整保留在 ocr_text
        assert "这是正文第一段" in (e.ocr_text or "")
        assert "这是正文第二段" in (e.ocr_text or "")
        # 图文流按文档顺序进 content_description：段落 1 → 图片 → 段落 3
        cd = e.content_description or ""
        assert "[段落 1]" in cd and "[段落 3]" in cd
        assert "[图片 1]" in cd and "image1" in cd
        assert r.failed == 0

    def test_docx_flow_items_order(self, tmp_path):
        """流式解析还原图文顺序：段落 → 图片 → 表格。"""
        import io
        from PIL import Image as PILImage
        import docx as docxlib
        from naskb.common.batch import _docx_flow_items

        p = tmp_path / "flow.docx"
        d = docxlib.Document()
        d.add_paragraph("第一段文字")
        buf = io.BytesIO()
        PILImage.new("RGB", (32, 32)).save(buf, "PNG")
        buf.seek(0)
        d.add_picture(buf)
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "甲"
        t.cell(0, 1).text = "乙"
        d.save(str(p))

        items = _docx_flow_items(str(p))
        kinds = [it["kind"] for it in items]
        assert kinds == ["text", "image", "table"], kinds
        assert items[0]["content"] == "第一段文字"
        assert items[1]["media"].endswith(".png")
        assert "段落" in items[1]["anchor"] and "inline" in items[1]["anchor"]
        assert "甲" in items[2]["content"] and "乙" in items[2]["content"]

    def test_docx_to_pdf_uses_word_com(self, tmp_path):
        """_docx_to_pdf：Word COM 渲染 docx → PDF，返回同目录 pdf 路径。"""
        from naskb.common.batch import _docx_to_pdf

        calls = {}

        class FakeDoc:
            def __init__(self, path):
                calls["open"] = path

            def SaveAs2(self, pdf_path, FileFormat):
                calls["save"] = (pdf_path, FileFormat)
                with open(pdf_path, "wb") as f:
                    f.write(b"%PDF fake")

            def Close(self, save):
                calls["close"] = save

        class FakeApp:
            Documents = None

            def __init__(self):
                self.Documents = self

            def Open(self, path, ReadOnly):
                return FakeDoc(path)

        src = tmp_path / "a.docx"
        src.write_bytes(b"PK fake")
        pdf = _docx_to_pdf(FakeApp(), str(src))
        assert pdf == str(tmp_path / "a.pdf")
        assert calls["save"][1] == 17          # wdFormatPDF
        assert calls["close"] is False          # 不保存改动

    def test_docx_to_pdf_fails_gracefully(self, tmp_path):
        """Word COM 不可用 → 返回 None（档位 2 跳过）。"""
        from naskb.common.batch import _docx_to_pdf

        class BoomApp:
            class Documents:
                @staticmethod
                def Open(path, ReadOnly):
                    raise RuntimeError("no word")

        assert _docx_to_pdf(BoomApp(), str(tmp_path / "x.docx")) is None

    def test_llm_failure_degrades(self, tmp_path, work):
        """LLM 失败降级：不崩溃，条目仍有提取文本。"""
        _make_files(tmp_path)
        fs = LocalAdapter(str(tmp_path))
        store = NaskbStore(fs)

        class BoomLLM(FakeLLM):
            def complete_json(self, prompt):
                raise RuntimeError("mock 失败")

        r = analyze_tree(fs, store, work, ".", llm=True, workers=2,
                         clients={"text": BoomLLM(), "vision": BoomLLM(),
                                  "audio": BoomLLM()})
        assert r.llm_failed == 6
        # 提取文本仍在（future 前已写入）
        e = store.get_entry(str(tmp_path / "笔记0.md"))
        assert e is not None
        assert "第 0 篇" in (e.ocr_text or "")
